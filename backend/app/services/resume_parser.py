from __future__ import annotations

import io
import re
from dataclasses import asdict, dataclass
from typing import Any, Dict, List, Sequence

import PyPDF2  # type: ignore[import]
import docx  # type: ignore[import]

_SECTION_KEYWORDS: Dict[str, Sequence[str]] = {
    "experience": ["опыт работы", "experience", "работа", "employment", "карьера"],
    "education": ["образование", "education", "обучение", "учёба"],
    "skills": ["навыки", "skills", "компетенции", "технологии", "стек"],
    "contacts": ["контакты", "contacts", "телефон", "email", "почта"],
    "summary": ["о себе", "summary", "обо мне", "профиль", "about"],
}

_TECH_SKILLS: Sequence[str] = (
    "python",
    "javascript",
    "typescript",
    "java",
    "c#",
    "c++",
    "go",
    "rust",
    "php",
    "ruby",
    "swift",
    "kotlin",
    "react",
    "angular",
    "vue",
    "html",
    "css",
    "sass",
    "redux",
    "next.js",
    "webpack",
    "node.js",
    "express",
    "fastapi",
    "django",
    "flask",
    "spring",
    ".net",
    "rails",
    "postgresql",
    "mysql",
    "mongodb",
    "redis",
    "oracle",
    "sql server",
    "elasticsearch",
    "docker",
    "kubernetes",
    "aws",
    "azure",
    "gcp",
    "ci/cd",
    "jenkins",
    "gitlab",
    "terraform",
    "git",
    "rest",
    "graphql",
    "microservices",
    "linux",
    "agile",
    "scrum",
)

_LANGUAGE_KEYWORDS: Dict[str, Sequence[str]] = {
    "английский": ["английский", "english", "англ"],
    "русский": ["русский", "russian"],
    "немецкий": ["немецкий", "german", "deutsch"],
    "французский": ["французский", "french"],
    "испанский": ["испанский", "spanish"],
    "китайский": ["китайский", "chinese"],
}

_LANGUAGE_LEVELS_PATTERN = r"(a1|a2|b1|b2|c1|c2|native|fluent|intermediate|basic)"
_EMAIL_PATTERN = r"[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Za-z]{2,}"
_PHONE_PATTERNS: Sequence[str] = (
    r"\+?\d[\s\-()]{0,3}(?:\d[\s\-()]{0,3}){6,}",
    r"\(?\d{3}\)?[\s\-]?\d{3}[\s\-]?\d{4}",
)


@dataclass
class ParsedResume:
    full_text: str
    name: str
    email: str
    phone: str
    skills: List[str]
    experience_years: float
    education: List[str]
    experience: List[str]
    summary: str
    languages: List[str]


class ResumeParser:
    def parse(self, file_content: bytes, file_type: str) -> Dict[str, Any]:
        text = self._extract_text(file_content, file_type)
        if not text.strip():
            return {"error": "Не удалось извлечь текст из файла"}

        result = ParsedResume(
            full_text=text,
            name=self._extract_name(text),
            email=self._extract_email(text),
            phone=self._extract_phone(text),
            skills=self._extract_skills(text),
            experience_years=self._extract_experience_years(text),
            education=self._extract_education(text),
            experience=self._extract_experience_section(text),
            summary=self._extract_summary(text),
            languages=self._extract_languages(text),
        )
        return asdict(result)

    def _extract_text(self, content: bytes, file_type: str) -> str:
        try:
            if file_type.lower() == "pdf":
                return self._extract_pdf_text(content)
            if file_type.lower() == "docx":
                return self._extract_docx_text(content)
            if file_type.lower() == "txt":
                return content.decode("utf-8", errors="ignore")
            return content.decode("utf-8", errors="ignore")
        except Exception as exc:  # pragma: no cover - defensive
            print(f"Error extracting text: {exc}")
            return ""

    def _extract_pdf_text(self, content: bytes) -> str:
        try:
            pdf_reader = PyPDF2.PdfReader(io.BytesIO(content))
        except Exception:
            return ""
        pages: List[str] = []
        for page in pdf_reader.pages:
            try:
                extracted = page.extract_text() or ""
            except Exception:
                extracted = ""
            if extracted:
                pages.append(extracted)
        return "\n".join(pages)

    def _extract_docx_text(self, content: bytes) -> str:
        try:
            document = docx.Document(io.BytesIO(content))
        except Exception:
            return ""
        parts: List[str] = []
        for paragraph in document.paragraphs:
            if paragraph.text:
                parts.append(str(paragraph.text))
        for table in document.tables:
            for row in table.rows:
                row_text = " ".join(str(cell.text)
                                    for cell in row.cells if cell.text)
                if row_text:
                    parts.append(row_text)
        return "\n".join(parts)

    def _extract_name(self, text: str) -> str:
        for line in text.splitlines()[:5]:
            line_stripped = line.strip()
            if not line_stripped or "@" in line_stripped or "http" in line_stripped:
                continue
            words = line_stripped.split()
            if 2 <= len(words) <= 4 and len(line_stripped) < 60:
                return line_stripped
        return "Unknown"

    def _extract_email(self, text: str) -> str:
        emails = re.findall(_EMAIL_PATTERN, text)
        return emails[0] if emails else ""

    def _extract_phone(self, text: str) -> str:
        for pattern in _PHONE_PATTERNS:
            phones = re.findall(pattern, text)
            if phones:
                return phones[0]
        return ""

    def _extract_skills(self, text: str) -> List[str]:
        text_lower = text.lower()
        found: List[str] = []
        for skill in _TECH_SKILLS:
            if skill in text_lower:
                found.append(skill.title())
        return sorted(set(found))

    def _extract_experience_years(self, text: str) -> float:
        matches = re.findall(
            r"(\d+)[\s\-]*(?:год|years?)", text, flags=re.IGNORECASE)
        if not matches:
            return 0.0
        values = [float(match) for match in matches]
        return max(values)

    def _extract_education(self, text: str) -> List[str]:
        section = self._find_section(text, "education")
        lines = [line.strip() for line in section.splitlines() if line.strip()]
        education: List[str] = []
        for line in lines:
            if re.search(r"университет|college|institute|academy", line, re.IGNORECASE):
                education.append(line)
        return education[:5]

    def _extract_experience_section(self, text: str) -> List[str]:
        section = self._find_section(text, "experience")
        if not section:
            return []
        entries: List[str] = []
        current = ""
        for line in section.splitlines():
            stripped = line.strip()
            if not stripped:
                continue
            if re.search(r"20\d{2}", stripped):
                if current:
                    entries.append(current)
                current = stripped
            else:
                current = f"{current} {stripped}".strip()
        if current:
            entries.append(current)
        return entries[:5]

    def _extract_summary(self, text: str) -> str:
        section = self._find_section(text, "summary")
        if section:
            return section[:500]
        lines = [line.strip()
                 for line in text.splitlines()[:20] if line.strip()]
        summary: List[str] = []
        for line in lines:
            if "@" in line:
                continue
            summary.append(line)
            if len(" ".join(summary)) > 300:
                break
        return " ".join(summary)[:500]

    def _extract_languages(self, text: str) -> List[str]:
        text_lower = text.lower()
        languages: List[str] = []
        for language, keywords in _LANGUAGE_KEYWORDS.items():
            for keyword in keywords:
                if keyword in text_lower:
                    level_match = re.search(
                        rf"{keyword}.*?{_LANGUAGE_LEVELS_PATTERN}", text_lower
                    )
                    if level_match:
                        languages.append(
                            f"{language} ({level_match.group(1)})")
                    else:
                        languages.append(language)
                    break
        return languages

    def _find_section(self, text: str, section_type: str) -> str:
        keywords = _SECTION_KEYWORDS.get(section_type, [])
        if not keywords:
            return ""
        lower_text = text.lower()
        for keyword in keywords:
            position = lower_text.find(keyword)
            if position == -1:
                continue
            section_text = text[position:]
            end = len(section_text)
            for other_type, other_keywords in _SECTION_KEYWORDS.items():
                if other_type == section_type:
                    continue
                for other_keyword in other_keywords:
                    other_pos = section_text.lower().find(other_keyword)
                    if 100 < other_pos < end:
                        end = other_pos
            return section_text[:end]
        return ""
