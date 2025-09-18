from __future__ import annotations

import io
import re
from pathlib import Path
from typing import Any, BinaryIO, Dict, IO, Mapping, Union, cast

from pdfminer.high_level import extract_text as pdf_extract_text
from docx import Document  # type: ignore[reportMissingTypeStubs]


def _normalize_text(text: str) -> str:
    normalized = text.replace("\r\n", "\n").replace("\r", "\n")
    normalized = normalized.replace("\u00A0", " ").replace("\u202F", " ")
    normalized = re.sub(r"[ \t]+", " ", normalized)
    normalized = re.sub(r"\n[ \t]+", "\n", normalized)
    normalized = re.sub(r"\n{3,}", "\n\n", normalized)
    return normalized.strip()


def _is_pdf_bytes(blob: bytes) -> bool:
    return blob[:4] == b"%PDF"


def _docx_text_from_bytes(blob: bytes) -> str:
    document = Document(io.BytesIO(blob))
    parts: list[str] = []
    for paragraph in document.paragraphs:
        if paragraph.text:
            parts.append(str(paragraph.text))
    for table in document.tables:
        for row in table.rows:
            cell_text = " ".join(str(cell.text) for cell in row.cells if cell.text)
            if cell_text:
                parts.append(cell_text)
    return "\n".join(parts)


def _pdf_text_from_bytes(blob: bytes, *, max_pages: int | None) -> str:
    if max_pages is None:
        return pdf_extract_text(io.BytesIO(blob))
    return pdf_extract_text(io.BytesIO(blob), maxpages=max_pages)


def _pdf_text_from_path(path: Path, *, max_pages: int | None) -> str:
    if max_pages is None:
        return pdf_extract_text(str(path))
    return pdf_extract_text(str(path), maxpages=max_pages)


def extract_text(src: Union[str, Path, bytes, IO[bytes]], max_pages: int | None = None) -> str:
    if isinstance(src, (str, Path)):
        path = Path(src)
        if path.suffix.lower() == ".docx":
            return _normalize_text(_docx_text_from_bytes(path.read_bytes()))
        text = _pdf_text_from_path(path, max_pages=max_pages)
        return _normalize_text(text)

    if hasattr(src, "read"):
        reader = cast(BinaryIO, src)
        data = reader.read()
    elif isinstance(src, bytes):
        data = src
    else:
        raise TypeError("Unsupported src type")

    if _is_pdf_bytes(data):
        text = _pdf_text_from_bytes(data, max_pages=max_pages)
        return _normalize_text(text)
    text = _docx_text_from_bytes(data)
    return _normalize_text(text)


def parse_resume(src: Union[str, Path, bytes, IO[bytes]]) -> Dict[str, Any]:
    text = extract_text(src)

    emails = re.findall(r"[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Za-z]{2,}", text)
    phones = re.findall(r"(?:\+?\d[\s\-()]{0,3}){7,}\d", text)

    skill_words = [
        "python",
        "sql",
        "pandas",
        "excel",
        "postgres",
        "postgresql",
        "oracle",
        "linux",
        "docker",
        "kubernetes",
        "spark",
        "hadoop",
        "airflow",
        "java",
        "c#",
        "javascript",
        "typescript",
        "react",
    ]
    skills = sorted(
        {
            word
            for word in skill_words
            if re.search(rf"\b{re.escape(word)}\b", text, re.IGNORECASE)
        }
    )

    language_patterns: Mapping[str, str] = {
        "english": r"\b(английский|english|англ)\b",
        "german": r"\b(немецкий|german)\b",
        "french": r"\b(французский|french)\b",
    }
    languages = [
        name for name, pattern in language_patterns.items() if re.search(pattern, text, re.IGNORECASE)
    ]

    contacts: Dict[str, str] = {}
    if emails:
        contacts["email"] = emails[0]
    if phones:
        contacts["phone"] = phones[0]

    return {
        "text": text,
        "contacts": contacts,
        "skills": skills,
        "languages": languages,
    }